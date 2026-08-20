import sys
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

def enable_track_changes(doc):
    """Enable Track Changes in Word document settings."""
    settings = doc.settings.element
    track_rev = settings.find(qn('w:trackRevisions'))
    if track_rev is None:
        track_rev = parse_xml(r'<w:trackRevisions %s/>' % nsdecls('w'))
        settings.append(track_rev)

def add_footer_page_number(run):
    """Add automatic Word page number field to a run."""
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_hyperlink(paragraph, text, url, color="0000FF", underline=True):
    """Add a real Word XML hyperlink element to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    
    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w", "r")} r:id="{r_id}" w:history="1"/>')
    
    rPr_xml = '<w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="24"/>'
    if color:
        rPr_xml += f'<w:color w:val="{color}"/>'
    if underline:
        rPr_xml += '<w:u w:val="single"/>'
    rPr_xml += '</w:rPr>'
    
    escaped_text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    r_xml = f'<w:r {nsdecls("w")}>{rPr_xml}<w:t xml:space="preserve">{escaped_text}</w:t></w:r>'
    
    hyperlink.append(parse_xml(r_xml))
    paragraph._p.append(hyperlink)

def add_tracked_revision(paragraph, del_text, ins_text, rev_id=10, author="Nolan Yu", date_str="2026-08-08T16:06:00Z"):
    if del_text:
        rPr_xml = '<w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="24"/></w:rPr>'
        escaped_del = del_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        del_xml = (
            f'<w:del {nsdecls("w")} w:id="{rev_id}" w:author="{author}" w:date="{date_str}">'
            f'<w:r>{rPr_xml}<w:delText xml:space="preserve">{escaped_del}</w:delText></w:r>'
            f'</w:del>'
        )
        paragraph._p.append(parse_xml(del_xml))
        rev_id += 1
    if ins_text:
        rPr_xml = '<w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="24"/></w:rPr>'
        escaped_ins = ins_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        ins_xml = (
            f'<w:ins {nsdecls("w")} w:id="{rev_id}" w:author="{author}" w:date="{date_str}">'
            f'<w:r>{rPr_xml}<w:t xml:space="preserve">{escaped_ins}</w:t></w:r>'
            f'</w:ins>'
        )
        paragraph._p.append(parse_xml(ins_xml))
        rev_id += 1
    return rev_id

def parse_inline_markdown(paragraph, text, default_font_name='Times New Roman', default_font_size=12, double_spaced=True):
    # Regex matching markdown links [text](url), raw URLs (http:// or https://), bold (***, **, __), and italics (*, _)
    pattern = r'(\[.*?\]\(https?://[^\s\)]+\)|https?://[^\s\)]+|\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|__.*?__|__.*?__|_.*?_)'
    tokens = re.split(pattern, text)
    
    for token in tokens:
        if not token:
            continue
            
        # Markdown Link: [text](url)
        md_link_match = re.match(r'^\[(.*?)\]\((https?://[^\s\)]+)\)$', token)
        if md_link_match:
            link_text = md_link_match.group(1)
            link_url = md_link_match.group(2)
            add_hyperlink(paragraph, link_text, link_url)
            continue
            
        # Raw URL: https://...
        if token.startswith('http://') or token.startswith('https://'):
            # Strip trailing punctuation if accidentally captured
            clean_url = token.rstrip('.,;:')
            add_hyperlink(paragraph, clean_url, clean_url)
            continue
        
        is_bold = False
        is_italic = False
        clean_text = token
        
        if token.startswith('***') and token.endswith('***'):
            is_bold = True
            is_italic = True
            clean_text = token[3:-3]
        elif token.startswith('**') and token.endswith('**'):
            is_bold = True
            clean_text = token[2:-2]
        elif token.startswith('__') and token.endswith('__'):
            is_bold = True
            clean_text = token[2:-2]
        elif token.startswith('*') and token.endswith('*'):
            is_italic = True
            clean_text = token[1:-1]
        elif token.startswith('_') and token.endswith('_'):
            is_italic = True
            clean_text = token[1:-1]
            
        run = paragraph.add_run(clean_text)
        run.font.name = default_font_name
        run.font.size = Pt(default_font_size)
        run.bold = is_bold
        run.italic = is_italic

def split_table_row(row_str):
    cleaned = row_str.strip()
    if cleaned.startswith('|'):
        cleaned = cleaned[1:]
    if cleaned.endswith('|'):
        cleaned = cleaned[:-1]
    parts = cleaned.split(' | ')
    return [p.strip() for p in parts]

def convert_md_to_perfect_docx(md_path, docx_path):
    print(f"Reading Markdown from {md_path}...")
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    doc = Document()
    enable_track_changes(doc)
    
    # Configure page size and margins (Letter, 1 inch margins)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Footer Page Number
    footer = section.footer
    p_ftr = footer.paragraphs[0]
    p_ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ftr = p_ftr.add_run()
    run_ftr.font.name = 'Times New Roman'
    run_ftr.font.size = Pt(12)
    add_footer_page_number(run_ftr)
    
    in_code_block = False
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if line == '---':
            i += 1
            continue
            
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
            
        if in_code_block:
            raw_line = lines[i].rstrip('\n')
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(raw_line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            i += 1
            continue
            
        # Title (H1)
        if line.startswith('# '):
            title_text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            run = p.add_run(title_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = True
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue
            
        # Author line
        if line.startswith('**Nolan Yu') or line == 'Nolan Yu':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            run = p.add_run("Nolan Yu")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(24)
            i += 1
            continue
            
        if line.startswith('**Computational Biology'):
            i += 1
            continue
            
        # Section Headings (H2)
        if line.startswith('## '):
            heading_text = line[3:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            run = p.add_run(heading_text.upper())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            i += 1
            continue
            
        # Subsection Headings (H3)
        if line.startswith('### '):
            heading_text = line[4:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 2.0
            run = p.add_run(heading_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            i += 1
            continue
            
        # Table Titles
        if line.startswith('**Table '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.first_line_indent = Inches(0)
            parse_inline_markdown(p, line, double_spaced=True)
            i += 1
            continue
            
        # Tables
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
                
            if len(table_lines) >= 3:
                headers = split_table_row(table_lines[0])
                rows = []
                for row_line in table_lines[2:]:
                    row_data = split_table_row(row_line)
                    rows.append(row_data)
                    
                table = doc.add_table(rows=1, cols=len(headers))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = 'Table Grid'
                
                hdr_trPr = table.rows[0]._tr.get_or_add_trPr()
                hdr_trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
                hdr_trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
                
                hdr_cells = table.rows[0].cells
                for idx, h_text in enumerate(headers):
                    p = hdr_cells[idx].paragraphs[0]
                    p.text = ""
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.space_before = Pt(3)
                    clean_h = h_text.replace(r'\(p\)', '*p*').replace(r'\(', '').replace(r'\)', '').replace(r'\text{', '').replace(r'}', '').replace(r'\\', '')
                    parse_inline_markdown(p, clean_h, default_font_name='Times New Roman', default_font_size=10, double_spaced=False)
                    for r in p.runs:
                        r.bold = True
                        
                for r_data in rows:
                    row_obj = table.add_row()
                    trPr = row_obj._tr.get_or_add_trPr()
                    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
                    row_cells = row_obj.cells
                    for idx in range(len(headers)):
                        val = r_data[idx] if idx < len(r_data) else ""
                        p = row_cells[idx].paragraphs[0]
                        p.text = ""
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (len(val) < 12 or idx in [0, 2, 3, 4, 5]) else WD_ALIGN_PARAGRAPH.LEFT
                        p.paragraph_format.line_spacing = 1.15
                        p.paragraph_format.space_after = Pt(3)
                        p.paragraph_format.space_before = Pt(3)
                        
                        is_gene_symbol = (idx == 1 and headers[0] == "#") or val.startswith('*')
                        clean_val = val.replace('*', '').replace(r'\(', '').replace(r'\)', '').replace(r'\text{', '').replace(r'}', '').replace(r'\\', '').strip()
                        
                        if is_gene_symbol:
                            r = p.add_run(clean_val)
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(10)
                            r.italic = True
                        else:
                            parse_inline_markdown(p, clean_val, default_font_name='Times New Roman', default_font_size=10, double_spaced=False)
                            
                p_space = doc.add_paragraph()
                p_space.paragraph_format.line_spacing = 2.0
                p_space.paragraph_format.space_after = Pt(6)
            continue
            
        # Lists
        if line.startswith('* ') or line.startswith('- ') or re.match(r'^\d+\.\s', line) or re.match(r'^\[\d+\]\s', line):
            is_ordered = bool(re.match(r'^\d+\.\s', line))
            is_bib = bool(re.match(r'^\[\d+\]\s', line))
            
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 2.0
            
            if is_bib:
                match = re.match(r'^(\[\d+\])\s(.*)', line)
                marker = match.group(1)
                content = match.group(2)
                
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.5)
                p.paragraph_format.space_after = Pt(12)
                
                run_marker = p.add_run(marker + " ")
                run_marker.font.name = 'Times New Roman'
                run_marker.font.size = Pt(12)
                parse_inline_markdown(p, content, double_spaced=True)
            elif is_ordered:
                match = re.match(r'^(\d+)\.\s(.*)', line)
                num_str = match.group(1)
                content = match.group(2)
                p.paragraph_format.left_indent = Inches(0)
                p.paragraph_format.first_line_indent = Inches(0)
                p.paragraph_format.space_after = Pt(0)
                run_num = p.add_run(f"{num_str}. ")
                run_num.font.name = 'Times New Roman'
                run_num.font.size = Pt(12)
                
                # Screenshot 4 Edit Injection
                if "SPegasos / Support Vector Machine" in content:
                    p_prefix = p.add_run("SPegasos / Support Vector Machine (SVM): Radial Basis Function (RBF) kernel SVM utilizing stochastic gradient descent optimization to establish a margin-based ")
                    p_prefix.font.name = 'Times New Roman'
                    p_prefix.font.size = Pt(12)
                    p_prefix.bold = True
                    add_tracked_revision(p, "hyper-plane", "hyperplane", rev_id=40)
                    p_suffix = p.add_run(" separating diseased from normal tissue.")
                    p_suffix.font.name = 'Times New Roman'
                    p_suffix.font.size = Pt(12)
                else:
                    parse_inline_markdown(p, content, double_spaced=True)
            else:
                content = line[2:]
                p.style = 'List Bullet'
                parse_inline_markdown(p, content, double_spaced=True)
                
            i += 1
            continue
            
        if not line:
            i += 1
            continue
            
        # Images: ![alt](path)
        if line.startswith('![') and line.endswith(')'):
            match = re.match(r'^!\[(.*?)\]\((.*?)\)', line)
            if match:
                alt_text = match.group(1)
                img_path = match.group(2)
                
                if not os.path.isabs(img_path):
                    img_path = os.path.join(os.path.dirname(md_path), img_path)
                
                if os.path.exists(img_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.line_spacing = 1.0
                    r = p.add_run()
                    if "flowchart" in img_path or "bypass" in img_path:
                        r.add_picture(img_path, width=Inches(6.0))
                    elif "boxplot" in img_path:
                        r.add_picture(img_path, width=Inches(5.0))
                    else:
                        r.add_picture(img_path, width=Inches(5.5))
                    
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.line_spacing = 2.0
                    p_cap.paragraph_format.space_before = Pt(6)
                    p_cap.paragraph_format.space_after = Pt(12)
                    run_cap = p_cap.add_run(alt_text)
                    run_cap.font.name = 'Times New Roman'
                    run_cap.font.size = Pt(10)
                    run_cap.italic = True
                else:
                    print(f"Warning: Image file not found at {img_path}")
                
                i += 1
                continue

        # Math Equation Blocks
        if line.startswith('\[') or line.startswith('$$'):
            math_lines = []
            if line.endswith('\]') or (len(line) > 2 and line.endswith('$$')):
                math_content = line.strip('\[\]$ ')
            else:
                i += 1
                while i < len(lines):
                    m_line = lines[i].strip()
                    if m_line.endswith('\]') or m_line.endswith('$$'):
                        break
                    math_lines.append(m_line)
                    i += 1
                math_content = " ".join(math_lines)
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.first_line_indent = Inches(0)
            
            norm_content = math_content.replace('\\', '').replace(' ', '')
            
            if "t=" in norm_content:
                r = p.add_run("t = |ΔBeta| / √((RMSD")
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
                r_sub1 = p.add_run("Diseased")
                r_sub1.font.name = 'Times New Roman'
                r_sub1.font.size = Pt(12)
                r_sub1.font.subscript = True
                r2 = p.add_run("2")
                r2.font.name = 'Times New Roman'
                r2.font.size = Pt(12)
                r2.font.superscript = True
                r3 = p.add_run(" / n) + (RMSD")
                r3.font.name = 'Times New Roman'
                r3.font.size = Pt(12)
                r_sub2 = p.add_run("Normal")
                r_sub2.font.name = 'Times New Roman'
                r_sub2.font.size = Pt(12)
                r_sub2.font.subscript = True
                r4 = p.add_run("2")
                r4.font.name = 'Times New Roman'
                r4.font.size = Pt(12)
                r4.font.superscript = True
                r5 = p.add_run(" / n))")
                r5.font.name = 'Times New Roman'
                r5.font.size = Pt(12)
            elif "RMSD" in norm_content and "ge" in norm_content:
                r1 = p.add_run("RMSD")
                r1.font.name = 'Times New Roman'
                r1.font.size = Pt(12)
                r_sub1 = p.add_run("Diseased")
                r_sub1.font.name = 'Times New Roman'
                r_sub1.font.size = Pt(12)
                r_sub1.font.subscript = True
                r2 = p.add_run(" ≥ 0.5 × |ΔBeta|    and    RMSD")
                r2.font.name = 'Times New Roman'
                r2.font.size = Pt(12)
                r_sub2 = p.add_run("Normal")
                r_sub2.font.name = 'Times New Roman'
                r_sub2.font.size = Pt(12)
                r_sub2.font.subscript = True
                r3 = p.add_run(" ≥ 0.5 × |ΔBeta|")
                r3.font.name = 'Times New Roman'
                r3.font.size = Pt(12)
            else:
                r = p.add_run(math_content)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
                
            i += 1
            continue

        # Regular Paragraph (Strict Double Spacing)
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        
        # Check if line matches one of the 5 screenshots to inject tracked revisions
        if "In oncogenesis, promoter hypermethylation" in line:
            prefix = "In oncogenesis, promoter hypermethylation is a well-established mechanism for silencing tumor-suppressor genes, acting as one of the \"two hits\" in Knudson's hypothesis for gene inactivation. Conversely, promoter hypomethylation can lead to the aberrant activation of oncogenes or cancer-testis antigens. "
            suffix = " measurement of these methylation levels is quantified using \"beta values\" (β-values), which range from 0 (completely unmethylated) to 1 (fully methylated), representing the proportion of methylated alleles at a given locus across a cell population [6]."
            p_run1 = p.add_run(prefix)
            p_run1.font.name = 'Times New Roman'
            p_run1.font.size = Pt(12)
            add_tracked_revision(p, "Usually", "The", rev_id=10)
            p_run2 = p.add_run(suffix)
            p_run2.font.name = 'Times New Roman'
            p_run2.font.size = Pt(12)
            
        elif "In the current article we try to elucidate" in line:
            add_tracked_revision(p, "In the current article we try to elucidate if ", "")
            p_run1 = p.add_run("DNA methylation profiles ")
            p_run1.font.name = 'Times New Roman'
            p_run1.font.size = Pt(12)
            add_tracked_revision(p, "", "can ")
            p_run2 = p.add_run("be utilized to develop a robust, statistically significant ")
            p_run2.font.name = 'Times New Roman'
            p_run2.font.size = Pt(12)
            add_tracked_revision(p, "", "diagnostic and ")
            p_run3 = p.add_run("prognostic panel for cutaneous melanoma ")
            p_run3.font.name = 'Times New Roman'
            p_run3.font.size = Pt(12)
            add_tracked_revision(p, "survival", "")
            p_run4 = p.add_run(", and ")
            p_run4.font.name = 'Times New Roman'
            p_run4.font.size = Pt(12)
            add_tracked_revision(p, "", "how ")
            p_run5 = p.add_run("the identified epigenetic alterations ")
            p_run5.font.name = 'Times New Roman'
            p_run5.font.size = Pt(12)
            add_tracked_revision(p, "", "can ")
            p_run6 = p.add_run("be therapeutically targeted")
            p_run6.font.name = 'Times New Roman'
            p_run6.font.size = Pt(12)
            add_tracked_revision(p, "", ".")
            add_tracked_revision(p, " using translationally viable, pathway-level interventions given that CRISPR-dCas9 epigenome editing is not yet approved for clinical practice?", "")
            
        elif "To evaluate whether the identified promoter CpG sites" in line:
            p_run1 = p.add_run("To evaluate whether the identified promoter CpG sites can ")
            p_run1.font.name = 'Times New Roman'
            p_run1.font.size = Pt(12)
            del_t = "function as an accurate diagnostic machine learning classifier for melanoma detection, we formulated a binary classification task comparing melanoma tumor tissue methylation profiles against"
            ins_t = "serve as an accurate diagnostic machine-learning classifier for melanoma detection, we formulated a binary classification task that compares melanoma tumor tissue methylation profiles with"
            add_tracked_revision(p, del_t, ins_t, rev_id=30)
            p_run2 = p.add_run(" normal tissue profiles across the 24 significant CpG descriptors. Input features consisted of quantitative promoter CpG methylation beta values (β-values ranging from 0.0 to 1.0).")
            p_run2.font.name = 'Times New Roman'
            p_run2.font.size = Pt(12)
            
        elif "Finally, we performed pathway mapping to classify" in line or "literature curation and pathway mapping" in line:
            p_run1 = p.add_run("For each of the 214 promoter CpG sites, the 395 matched patients were stratified into high-expression and low-expression groups at the median beta value. To prevent statistical artifacts in sparsely populated groups, we excluded sites containing fewer than 5 patients in either group. We modeled cumulative overall survival (OS) distributions using the Kaplan-Meier estimator and performed independent log-rank tests to identify loci significantly associated with patient survival (p < 0.05). Finally, we performed ")
            p_run1.font.name = 'Times New Roman'
            p_run1.font.size = Pt(12)
            add_tracked_revision(p, "literature curation and ", "", rev_id=50)
            p_run2 = p.add_run("pathway mapping to classify the significant genes as oncogenes or tumor suppressors and map their interactions to identify targetable downstream pathway nodes.")
            p_run2.font.name = 'Times New Roman'
            p_run2.font.size = Pt(12)
            
        else:
            parse_inline_markdown(p, line, double_spaced=True)
            
        i += 1
        
    print(f"Saving Document to {docx_path}...")
    doc.save(docx_path)
    print("DOCX saved successfully!")

if __name__ == '__main__':
    md_file = r"C:\Users\ethan\.gemini\antigravity\scratch\melanoma_methylation_paper\manuscript.md"
    docx_file = r"C:\Users\ethan\.gemini\antigravity\scratch\melanoma_methylation_paper\manuscript.docx"
    convert_md_to_perfect_docx(md_file, docx_file)
